"""DOCX utility: read paragraphs/tables and optionally append text via CLI.

Requires `python-docx` (install with: pip install python-docx).
"""

from __future__ import annotations

import argparse
import logging
import sys
from pathlib import Path
from typing import Any, List, Optional

import docx  # python-docx

logger = logging.getLogger(__name__)


class DocxError(RuntimeError):
    """Raised when DOCX operations fail."""


def load_document(input_path: Optional[Path]) -> Any:
    """Load an existing DOCX or create a new one if `input_path` is None."""
    try:
        if input_path is None:
            return docx.Document()
        return docx.Document(str(input_path))
    except Exception as ex:  # noqa: BLE001
        raise DocxError(f"Failed to load document: {ex}") from ex


def get_paragraphs_text(document: Any) -> List[str]:
    return [p.text for p in document.paragraphs]


def get_tables_text(document: Any) -> List[List[List[str]]]:
    """Return tables as a list of tables -> rows -> cell texts."""
    tables: List[List[List[str]]] = []
    for table in document.tables:
        table_rows: List[List[str]] = []
        for row in table.rows:
            table_rows.append([cell.text for cell in row.cells])
        tables.append(table_rows)
    return tables


def append_paragraph(document: Any, text: str) -> None:
    document.add_paragraph(text)


def save_document(document: Any, output_path: Path) -> None:
    try:
        document.save(str(output_path))
    except Exception as ex:  # noqa: BLE001
        raise DocxError(f"Failed to save document: {ex}") from ex


def parse_arguments() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Read/modify DOCX files.",
        formatter_class=argparse.ArgumentDefaultsHelpFormatter,
    )
    parser.add_argument("--input", type=Path, help="Path to input .docx (optional)")
    parser.add_argument(
        "--output", type=Path, default=Path("output.docx"), help="Path to save .docx"
    )
    parser.add_argument(
        "--append-text", type=str, help="Append a paragraph to the document"
    )
    parser.add_argument(
        "--print-paragraphs", action="store_true", help="Print paragraphs to stdout"
    )
    parser.add_argument(
        "--print-tables", action="store_true", help="Print table cell texts to stdout"
    )
    parser.add_argument(
        "--log-level",
        choices=["CRITICAL", "ERROR", "WARNING", "INFO", "DEBUG"],
        default="INFO",
        help="Logging verbosity",
    )
    return parser.parse_args()


def main() -> int:
    args = parse_arguments()

    logging.basicConfig(
        level=getattr(logging, args.log_level.upper(), logging.INFO),
        format="%(message)s",
        stream=sys.stdout,
    )

    try:
        document = load_document(args.input)
    except DocxError as ex:
        logger.error(str(ex))
        return 2

    if args.print_paragraphs:
        paras = get_paragraphs_text(document)
        for line in paras:
            print(line)

    if args.print_tables:
        tables = get_tables_text(document)
        for t_idx, table in enumerate(tables):
            logger.info(f"Table {t_idx}")
            for r_idx, row in enumerate(table):
                logger.info("\t" + " | ".join(row))

    if args.append_text:
        append_paragraph(document, args.append_text)
        logger.info("Appended paragraph to document.")

    try:
        save_document(document, args.output)
    except DocxError as ex:
        logger.error(str(ex))
        return 1

    logger.info(f"Saved to {args.output}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
